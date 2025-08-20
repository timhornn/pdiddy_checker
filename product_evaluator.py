import easyocr
import cv2
import numpy as np
import re
import json
import requests
import docx
from collections import defaultdict
import os
import logging
import time

# Настройка логирования
logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        level=logging.INFO
    )


class ProductEvaluator:
    def __init__(self):
        """
        Инициализация оценщика
        """
        self.qwen_api_key = "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"

        # Инициализация EasyOCR (только русский язык)
        logger.info("Загрузка модели EasyOCR (русский язык)...")
        try:
            self.reader = easyocr.Reader(['ru'], gpu=False)
            logger.info("Модель EasyOCR успешно загружена.")
        except Exception as e:
            logger.error(f"Ошибка при загрузке EasyOCR: {str(e)}")
            raise

    def preprocess_image(self, image_path):
        """
        Подготавливает изображение для OCR с улучшенной обработкой
        """
        logger.info(f"Начало обработки изображения: {image_path}")

        # Загрузка изображения
        img = cv2.imread(image_path)
        if img is None:
            raise Exception(f"Не удалось загрузить изображение: {image_path}")

        # Конвертация в оттенки серого
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        logger.debug("Изображение преобразовано в оттенки серого")

        # Удаление шума
        denoised = cv2.fastNlMeansDenoising(gray, h=10)
        logger.debug("Удален шум с помощью fastNlMeansDenoising")

        # Адаптивная бинаризация
        thresh = cv2.adaptiveThreshold(
            denoised, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 15, 5
        )
        logger.debug("Применена адаптивная бинаризация")

        # Увеличение размера изображения для улучшения распознавания
        scale_percent = 150
        width = int(thresh.shape[1] * scale_percent / 100)
        height = int(thresh.shape[0] * scale_percent / 100)
        dim = (width, height)
        resized = cv2.resize(thresh, dim, interpolation=cv2.INTER_CUBIC)
        logger.debug(f"Изображение увеличено до {scale_percent}%")

        return resized

    def extract_text_from_image(self, image_path):
        """
        Извлекает текст из изображения с помощью EasyOCR
        """
        logger.info("Начало распознавания текста с изображения")

        try:
            # Предварительная обработка
            processed_img = self.preprocess_image(image_path)

            # Распознавание текста
            logger.debug("Запуск распознавания текста через EasyOCR...")
            results = self.reader.readtext(
                processed_img,
                detail=0,
                paragraph=True,
                min_size=20,
                text_threshold=0.7,
                low_text=0.4
            )

            # Объединение в строку
            text = " ".join(results)
            logger.debug(f"Распознанный текст: {text}")

            # Удаление повторяющихся пробелов
            text = re.sub(r'\s+', ' ', text).strip()

            return text
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста: {str(e)}")
            raise

    def get_who_requirements(self, requirements_path='Инструкция_по_оценке_продукта_питания.docx'):
        """
        Извлекает требования ВОЗ из документа
        """
        logger.info(f"Извлечение требований ВОЗ из файла: {requirements_path}")

        # Проверка существования файла
        if not os.path.exists(requirements_path):
            raise FileNotFoundError(f"Файл {requirements_path} не найден")

        try:
            # Попытка открыть как .docx (даже если расширение .doc)
            doc = docx.Document(requirements_path)
        except Exception:
            # Если не получается, попробуем временно переименовать
            if requirements_path.endswith('.doc') and not requirements_path.endswith('.docx'):
                temp_path = requirements_path + 'x'
                try:
                    os.rename(requirements_path, temp_path)
                    doc = docx.Document(temp_path)
                    # Вернем обратно имя файла после обработки
                    os.rename(temp_path, requirements_path)
                except Exception:
                    raise ValueError(f"Файл {requirements_path} не является корректным документом Word")

        # Сбор всего текста документа
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Объединение текста
        text = "\n".join(full_text)

        # Удаление повторяющихся пробелов
        text = re.sub(r'\s+', ' ', text).strip()

        logger.info("Требования ВОЗ успешно извлечены")
        return text

    def analyze_with_qwen(self, who_requirements, product_text):
        """
        Отправляет запрос в Qwen3 API для анализа соответствия продукта требованиям ВОЗ
        """
        logger.info("Отправка запроса в Qwen3 API для анализа соответствия продукта требованиям ВОЗ")

        # Формируем промпт
        prompt = self._create_analysis_prompt(who_requirements, product_text)

        # Подготовка запроса к API
        headers = {
            "Authorization": f"Bearer {self.qwen_api_key}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "qwen/qwen3-235b-a22b",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.2,
            "max_tokens": 1000
        }

        try:
            # Отправка запроса
            logger.debug("Отправка запроса к Qwen3 API...")
            start_time = time.time()
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=data
            )
            elapsed = time.time() - start_time
            logger.debug(f"Получен ответ от API за {elapsed:.2f} секунд")

            # Проверка статуса ответа
            if response.status_code != 200:
                logger.warning(f"Ошибка API: {response.status_code} - {response.text}")
                return self._create_fallback_report(product_text)

            # Извлечение ответа
            result = response.json()
            analysis = result['choices'][0]['message']['content'].strip()

            logger.debug(f"Получен анализ от Qwen3: {analysis}")
            return analysis

        except Exception as e:
            logger.error(f"Ошибка при обращении к Qwen3 API: {str(e)}")
            return self._create_fallback_report(product_text)

    def _create_analysis_prompt(self, who_requirements, product_text):
        """
        Создает промпт для анализа соответствия продукта требованиям ВОЗ
        """
        prompt = (
            "Ты являешься экспертом по оценке продуктов детского питания в соответствии с требованиями ВОЗ. Тебе предоставлены:\n\n"

            "1. Требования ВОЗ по оценке продуктов питания для детей грудного и раннего возраста от 6 до 36 месяцев\n"
            "2. Текст этикетки продукта детского питания\n\n"

            "Твоя задача:\n"
            "1. Определить категорию продукта на основе текста этикетки\n"
            "2. Проверить соответствие продукта требованиям ВОЗ для этой категории\n"
            "3. Выявить все нарушения (если они есть)\n"
            "4. Сформировать отчет в следующем строгом формате:\n\n"

            "РЕЗУЛЬТАТ ОЦЕНКИ СООТВЕТСТВИЯ ПРОДУКТА ТРЕБОВАНИЯМ ВОЗ\n"
            "============================================================\n\n"

            "ОПРЕДЕЛЕННАЯ КАТЕГОРИЯ: [код категории] - [название категории]\n"
            "НАЗВАНИЕ ПРОДУКТА: [название продукта]\n"
            "ВОЗРАСТНАЯ МАРКИРОВКА: [возрастная маркировка]\n\n"

            "РЕЗУЛЬТАТ: [сообщение о результате]\n\n"

            "НАЙДЕННЫЕ НАРУШЕНИЯ:\n"
            "1. [нарушение 1]\n"
            "2. [нарушение 2]\n"
            "...\n\n"

            "РЕКОМЕНДАЦИИ:\n"
            "[рекомендации]\n\n"

            "Требования ВОЗ:\n\n"
            f"{who_requirements}\n\n"

            "Текст этикетки продукта:\n\n"
            f"{product_text}\n\n"

            "Пожалуйста, проведи анализ и сформируй отчет в указанном формате. "
            "Если категория продукта не может быть определена, укажи 'Неизвестная категория'. "
            "Если возрастная маркировка отсутствует, укажи 'Не указана'. "
            "Если нарушений нет, укажи 'Нарушения не обнаружены'."
        )

        return prompt

    def _create_fallback_report(self, product_text):
        """
        Создает резервный отчет в случае недоступности API
        """
        # Попробуем определить категорию и возрастную маркировку простым анализом
        category = "Неизвестная категория"
        age_label = "Не указана"
        product_name = "Не определено"

        # Определение категории
        if re.search(r'сок|напиток|нектар|смузи', product_text):
            category = "8 - Напитки"
        elif re.search(r'йогурт|молоко|сыр|творог|пудинг', product_text):
            category = "2 - Молочные продукты"
        elif re.search(r'каша|крупа|мюсли|рис|макароны', product_text):
            category = "1 - Сухие каши"
        elif re.search(r'пюре|фрукт|яблоко|груша|банан', product_text):
            category = "3.1 - Фруктовые и овощные пюре"

        # Определение возрастной маркировки
        age_match = re.search(r'(?:возраст|для|детей)\s*(?:от\s*)?(\d+)\s*(?:месяц|мес)', product_text)
        if age_match:
            age_label = f"С {age_match.group(1)} месяцев"

        # Определение названия продукта
        lines = product_text.split('\n')
        if lines:
            for line in lines[:3]:
                if re.search(r'[а-яА-Я]', line) and len(line) > 5:
                    product_name = line.strip().title()
                    break

        # Формирование отчета
        report = "РЕЗУЛЬТАТ ОЦЕНКИ СООТВЕТСТВИЯ ПРОДУКТА ТРЕБОВАНИЯМ ВОЗ\n"
        report += "=" * 60 + "\n\n"
        report += f"ОПРЕДЕЛЕННАЯ КАТЕГОРИЯ: {category}\n"
        report += f"НАЗВАНИЕ ПРОДУКТА: {product_name}\n"
        report += f"ВОЗРАСТНАЯ МАРКИРОВКА: {age_label}\n\n"
        report += "РЕЗУЛЬТАТ: Не удалось провести полный анализ из-за недоступности API\n\n"
        report += "РЕКОМЕНДАЦИИ:\n"
        report += "❌ Не удалось провести полный анализ. Пожалуйста, повторите попытку позже.\n"
        report += "Система временно использует упрощенный анализ на основе ключевых слов."

        return report

    def evaluate_product(self, image_path):
        """
        Полная оценка продукта с использованием Qwen3 API
        """
        logger.info("Начало оценки продукта...")

        # Извлечение текста из изображения
        product_text = self.extract_text_from_image(image_path)
        logger.info(f"Извлеченный текст с этикетки: {product_text[:200]}{'...' if len(product_text) > 200 else ''}")

        # Извлечение требований ВОЗ
        try:
            who_requirements = self.get_who_requirements()
            logger.info(f"Требования ВОЗ: {who_requirements[:200]}{'...' if len(who_requirements) > 200 else ''}")
        except Exception as e:
            logger.error(f"Ошибка при извлечении требований ВОЗ: {str(e)}")
            who_requirements = "Не удалось извлечь требования ВОЗ из документа"

        # Анализ с помощью Qwen3 API
        analysis = self.analyze_with_qwen(who_requirements, product_text)

        logger.info("Оценка продукта завершена")
        return analysis

    def generate_evaluation_report(self, analysis):
        """
        Возвращает отчет об оценке (уже сформированный Qwen3)
        """
        return analysis